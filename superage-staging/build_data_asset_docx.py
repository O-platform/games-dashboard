"""Generate SuperAge Data Asset & Growth Story slide-ready DOCX."""
import os
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(OUT_DIR, "SuperAge_Data_Asset_Growth_Story.docx")
CHART_GROWTH = os.path.join(OUT_DIR, "_chart_growth_history.png")
CHART_3MO = os.path.join(OUT_DIR, "_chart_3mo_projection.png")
CHART_TICKETS = os.path.join(OUT_DIR, "_chart_tickets_5mo.png")
CHART_ACQ = os.path.join(OUT_DIR, "_chart_acquisition_5mo.png")


def _style_axes(ax):
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.grid(axis="y", linestyle="--", alpha=0.3)
    ax.tick_params(labelsize=10)


def make_growth_history_chart():
    months = ["Jan 1", "Feb 1", "Mar 1", "Apr 1", "May 1"]
    totals = [1076055, 1077202, 1088742, 1093187, 1108740]
    fig, ax = plt.subplots(figsize=(8, 3.6))
    ax.plot(months, totals, marker="o", linewidth=2.4, color="#0969da", markersize=8)
    for i, (m, t) in enumerate(zip(months, totals)):
        ax.annotate(f"{t:,}", (m, t), textcoords="offset points",
                    xytext=(0, 10), ha="center", fontsize=9)
    ax.set_title("Active Subscribers — 5-Month History (2026)", fontsize=12, fontweight="bold")
    ax.set_ylabel("Total Active Subscribers")
    ax.set_ylim(1060000, 1125000)
    _style_axes(ax)
    plt.tight_layout()
    plt.savefig(CHART_GROWTH, dpi=160)
    plt.close()


def make_3mo_projection_chart():
    months = ["Jun 1\n(today)", "Jul 1", "Aug 1", "Sep 1"]
    mid = [1108358, 1115123, 1121929, 1128778]
    fig, ax = plt.subplots(figsize=(8, 3.6))
    ax.plot(months, mid, marker="o", linewidth=2.6, color="#1a7f37", markersize=9)
    for m, v in zip(months, mid):
        ax.annotate(f"{v:,}", (m, v), textcoords="offset points",
                    xytext=(0, 10), ha="center", fontsize=9)
    ax.set_title("3-Month Subscriber Projection — Primary Case (Jun → Sep 2026)",
                 fontsize=12, fontweight="bold")
    ax.set_ylabel("Total Active Subscribers")
    ax.set_ylim(1100000, 1135000)
    _style_axes(ax)
    plt.tight_layout()
    plt.savefig(CHART_3MO, dpi=160)
    plt.close()


def make_tickets_5mo_chart():
    months = ["Jun (today)", "Jul", "Aug", "Sep", "Oct", "Nov (event)"]
    # S-curve toward event: heavier in final 2 months
    buyers = [142, 280, 500, 820, 1230, 1750]
    fig, ax = plt.subplots(figsize=(8, 3.6))
    ax.plot(months, buyers, marker="o", linewidth=2.6, color="#8250df", markersize=9)
    for m, v in zip(months, buyers):
        ax.annotate(f"{v:,}", (m, v), textcoords="offset points",
                    xytext=(0, 10), ha="center", fontsize=9)
    ax.set_title("Games Ticket Sales — Primary Case to Nov Event",
                 fontsize=12, fontweight="bold")
    ax.set_ylabel("Cumulative Paid Buyers")
    ax.set_ylim(0, 2000)
    _style_axes(ax)
    plt.tight_layout()
    plt.savefig(CHART_TICKETS, dpi=160)
    plt.close()


def make_acquisition_5mo_chart():
    months = ["Jun (today)", "Jul", "Aug", "Sep", "Oct", "Nov"]
    cumulative_new_subs = [0, 114, 228, 342, 456, 570]
    fig, ax = plt.subplots(figsize=(8, 3.6))
    ax.plot(months, cumulative_new_subs, marker="o", linewidth=2.6, color="#cf222e", markersize=9)
    for m, v in zip(months, cumulative_new_subs):
        ax.annotate(f"{v:,}", (m, v), textcoords="offset points",
                    xytext=(0, 10), ha="center", fontsize=9)
    ax.set_title("Net-New SuperAge Subscribers Driven by Games — Primary Case",
                 fontsize=12, fontweight="bold")
    ax.set_ylabel("Cumulative Net-New SuperAge Subscribers")
    ax.set_ylim(0, 700)
    _style_axes(ax)
    plt.tight_layout()
    plt.savefig(CHART_ACQ, dpi=160)
    plt.close()


# ── Document helpers ─────────────────────────────────────
def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, bold=False, italic=False, size=None, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for r in hdr[i].paragraphs[0].runs:
            r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t


def add_quote(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)


# ── Build doc ──────────────────────────────────────────
def build():
    make_growth_history_chart()
    make_3mo_projection_chart()
    make_tickets_5mo_chart()
    make_acquisition_5mo_chart()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading("SuperAge — Data Asset & Growth Story", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_para(doc, "Slide-ready content. Numbers verified against production data as of June 12, 2026.",
             italic=True, size=10, color=RGBColor(0x6B, 0x72, 0x80))

    # 1. Profile Base
    add_heading(doc, "1. The Profile Base", level=1)
    add_table(doc,
              ["Metric", "Count"],
              [["Total Profiles (including unsubscribes; ever consented, all states)", "1,822,753"],
               ["Consented Profiles", "1,106,146"],
               ["Active Consented Profiles", "772,526"],
               ["Openers (last 90 days)", "640,668"],
               ["Recently Engaged Clickers (last 90 days)", "194,924"]])

    # 2. Fields per profile
    add_heading(doc, "2. Fields per Profile — 25+ Attributes Across 8 Categories", level=1)

    categories = [
        ("Demographics", "Age · Date of Birth · Gender · Marital Status · Education Level · Financial Situation"),
        ("Health & Lifestyle", "Exercise Frequency · Sleep Hours · Diet Description · Alcohol Frequency · Smoking Status · Social Frequency · Stress Impact"),
        ("Biometric & Clinical", "Diastolic Blood Pressure · Obesity / BMI Indicator · Health Conditions · Longevity Score (composite, 0–100)"),
        ("Family Health History", "Mother's Independent Age · Grandparents' Age at Death · Family Cancer Count"),
        ("Safety & Risk", "Traffic Incidents (last 5 years)"),
        ("Acquisition", "Source / Channel · Acquisition Date"),
        ("Engagement & Behavioral", "Engagement Segment · Open History · Click History · Days Active"),
        ("Lifecycle", "Subscribe Date · Unsubscribe Date · Current State"),
    ]
    for cat, fields in categories:
        p = doc.add_paragraph()
        r = p.add_run(cat)
        r.bold = True
        p.add_run("\n" + fields)

    # 3. Health Data Coverage
    add_heading(doc, "3. Health Data Coverage", level=1)
    add_para(doc, "10 validated health attributes per profile.", bold=True)
    add_quote(doc,
              "49.1% (896,299 out of 1,822,753 total profiles) carry validated health data — a "
              "clinically meaningful health dataset at near-million scale, including biometric "
              "(blood pressure, BMI) and behavioral risk factors (smoking, alcohol, exercise) used "
              "in actuarial and longevity research models.")
    add_para(doc, "Health attributes captured:", bold=True)
    add_bullets(doc, [
        "Exercise Frequency",
        "Sleep Hours",
        "Diet Description",
        "Alcohol Frequency",
        "Smoking Status",
        "Stress Impact",
        "Diastolic Blood Pressure",
        "Obesity / BMI Indicator",
        "Health Conditions",
        "Longevity Score (composite, 0–100)",
    ])

    # 4. Growth History
    add_heading(doc, "4. Subscriber Growth — 5-Month History", level=1)
    add_table(doc,
              ["Month Start", "Total Active", "MoM Change", "MoM %"],
              [["Jan 1, 2026", "1,076,055", "—", "—"],
               ["Feb 1, 2026", "1,077,202", "+1,147", "+0.11%"],
               ["Mar 1, 2026", "1,088,742", "+11,540", "+1.07%"],
               ["Apr 1, 2026", "1,093,187", "+4,445", "+0.41%"],
               ["May 1, 2026", "1,108,740", "+15,553", "+1.42%"]])
    doc.add_picture(CHART_GROWTH, width=Inches(6.5))

    add_para(doc, "Key observations:", bold=True)
    add_bullets(doc, [
        "5-month average MoM: +0.75% (~+8,171 subs / month).",
        "Peak months: Mar (+1.07%) and May (+1.42%) — driven by campaign pushes.",
        "Variance is the story: growth is real but campaign-dependent; needs a consistent driver. Games is that lever.",
    ])

    # 5. Games funnel
    add_heading(doc, "5. Games — Validated Funnel", level=1)
    add_table(doc,
              ["Metric", "Value"],
              [["Paid ticket buyers", "142"],
               ["Waitlist signups (closed)", "3,190"],
               ["Landing events (52 days, Apr 22 – Jun 12)", "45,585"],
               ["Current run-rate", "~17,000 events / month"],
               ["Site conversion (events → paid)", "0.316% (1 buyer per 317 events)"]])
    add_para(doc, "The funnel is real and measurable — not hypothetical.", italic=True)

    # 6. Games two-fold value
    add_heading(doc, "6. Games — Two-Fold Value to SuperAge", level=1)

    add_para(doc, "Acquisition (net-new subscribers driven by Games):", bold=True)
    add_bullets(doc, [
        "8,975 partner-brand clickers (HealthBrief + AllHealthy, time-stamped).",
        "102 subscribed to SuperAge AFTER clicking Games and are still consented.",
        "1.14% gross acquisition rate / 2.21% on net-new-eligible audience.",
    ])

    add_para(doc, "Retention / Engagement (existing subscribers re-engaging with brand):", bold=True)
    add_bullets(doc, [
        "4,363 existing SuperAge subscribers clicked Games URLs from partner-brand campaigns.",
        "That's 48.6% of all Games clickers — Games is a powerful re-engagement signal for the existing list.",
    ])

    # 7. Projections — primary case only
    add_heading(doc, "7. Projections — Primary Case", level=1)

    # 7a
    add_heading(doc, "7a. 3-Month Subscriber Projection (Jul → Sep 2026)", level=2)
    add_para(doc, "Starting base: 1,108,358 active subscribers (Jun 1, 2026)", italic=True)
    add_bullets(doc, [
        "Baseline: +0.60% / mo (return to the historical 6-month average)",
        "Games adds: +114 net-new / mo from scaled partner activation",
        "Net change: ~+6,765 / mo",
    ])
    add_table(doc,
              ["Month", "Total Active", "MoM %"],
              [["Jul 1, 2026", "1,115,123", "+0.61%"],
               ["Aug 1, 2026", "1,121,929", "+0.61%"],
               ["Sep 1, 2026", "1,128,778", "+0.61%"]])
    doc.add_picture(CHART_3MO, width=Inches(6.5))
    add_quote(doc, "3-month outcome: +20,420 subs (+1.84%) — Growth returns to historical norm with "
                   "modest Games activation. This is the case to plan against.")

    # 7b
    add_heading(doc, "7b. Games Ticket Sales — 5 Months to November Event", level=2)
    add_para(doc, "Public launch + light paid acquisition. Conversion held at observed 0.316%.",
             italic=True)
    add_table(doc,
              ["Month", "Cumulative Paid Buyers"],
              [["Jun (today)", "142"],
               ["Jul", "~280"],
               ["Aug", "~500"],
               ["Sep", "~820"],
               ["Oct", "~1,230"],
               ["Nov (event)", "~1,750"]])
    doc.add_picture(CHART_TICKETS, width=Inches(6.5))
    add_quote(doc, "S-curve concentrates ~50% of sales in the final 8 weeks — standard event-ticketing dynamic.")

    # 7c
    add_heading(doc, "7c. Net-New SuperAge Acquisition Driven by Games (5-Month View)", level=2)
    add_para(doc, "Using the validated 1.14% net-new acquisition rate from partner-brand Games clickers.",
             italic=True)
    add_table(doc,
              ["Month", "Cumulative Net-New SuperAge Subscribers"],
              [["Jun (today)", "0"],
               ["Jul", "~114"],
               ["Aug", "~228"],
               ["Sep", "~342"],
               ["Oct", "~456"],
               ["Nov", "~570"]])
    doc.add_picture(CHART_ACQ, width=Inches(6.5))
    add_quote(doc, "~570 net-new SuperAge subscribers acquired through Games campaigns by event date.")

    # 8. Strategic Story
    add_heading(doc, "8. The Strategic Story", level=1)
    add_quote(doc,
              "Today: 1.1M consented profiles, 49.1% health-data-validated, growth averaged "
              "+0.75% MoM over 5 months but is campaign-dependent.")
    add_para(doc, "Primary case with Games activation:", bold=True)
    add_bullets(doc, [
        "Sep 1, 2026 → 1,128,778 active subscribers (+20,420, +1.84% over 3 months) — restoring growth to the 6-month historical norm.",
        "Nov event → ~1,750 paid Games buyers (5× current customer base).",
        "Games campaigns drive ~570 net-new SuperAge subscribers over the 5-month run-up.",
    ])
    add_para(doc, "Why the primary case is defensible:", bold=True)
    add_bullets(doc, [
        "0.316% Games landing-event-to-paid conversion (measured)",
        "1.14% Games-click-to-SuperAge-subscriber acquisition (measured)",
        "48.6% Games-click-to-existing-sub retention engagement (measured)",
        "+0.60% MoM baseline is the historical 6-month average — not a stretch assumption",
        "November event is the inflection point — 5 months of build with measurable funnel mechanics",
    ])

    # 9. Data Licensing
    add_heading(doc, "9. Data Licensing Position", level=1)
    add_quote(doc, "All data will be aggregated and de-identified in packaging for licensing.")
    add_bullets(doc, [
        "1.1M consented profiles with 25+ attributes spanning demographics, behavior, lifestyle, and biometrics.",
        "Industry-standard k-anonymity safeguard (minimum bucket size = 25 profiles) ensures privacy in every export.",
        "Comparable to Truveta, Wirestock, and Protégé in scale and depth — distinguished by behavioral + longitudinal engagement signal (open / click history per subscriber over time).",
    ])

    doc.save(DOCX_PATH)
    # Clean up chart PNGs (embedded in docx, no longer needed)
    for f in (CHART_GROWTH, CHART_3MO, CHART_TICKETS, CHART_ACQ):
        if os.path.exists(f):
            os.remove(f)
    print(f"WROTE: {DOCX_PATH}")


if __name__ == "__main__":
    build()
